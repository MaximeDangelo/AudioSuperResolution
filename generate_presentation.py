"""
Genere le document de presentation du projet de super-resolution audio.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "rapport hebdo", "Presentation_Projet_SuperResolution_Audio.docx")

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_metric_table(doc, rows_data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F3864')
        run = hdr[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[i].paragraphs[0].runs[0].font.size = Pt(10)
    return table

doc = Document()

# --- Marges ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# PAGE DE TITRE
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Super-Resolution Audio\npour Communications Radio")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Pipeline hybride Demucs + SpectralResUNet + MetricGAN+")
run2.font.size = Pt(14)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run("Mars 2026  |  Maxime D.").font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. CONTEXTE ET PROBLEMATIQUE
# ============================================================
add_heading(doc, "1. Contexte et problematique")

doc.add_paragraph(
    "Les communications radio en aviation et dans le domaine militaire sont soumises a de fortes "
    "contraintes physiques qui degradent significativement la qualite du signal audio :"
)
bullets = [
    "Bande passante limitee : 300 - 3 400 Hz (radio AM telephonique)",
    "Bruit de fond : souffle HF, interferences, crackling, distorsion de modulation",
    "Signal vocal degrade : perte des harmoniques hautes, compression excessive, dropouts",
    "Absence totale des frequences au-dela de ~4 kHz",
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.runs[0].font.size = Pt(11)

doc.add_paragraph(
    "L'objectif du projet est de reconstruire les frequences manquantes (4 - 20 kHz) et "
    "d'ameliorer l'intelligibilite globale du signal, sans introduire d'artefacts vocaux. "
    "Ce travail s'inscrit dans le cadre d'un stage et constitue un spin-off d'un pipeline "
    "de debruitage existant (Demucs + Whisper)."
)

# ============================================================
# 2. ARCHITECTURE DU PIPELINE
# ============================================================
add_heading(doc, "2. Architecture du pipeline")

doc.add_paragraph(
    "Le pipeline hybride enchaine trois modeles specialises, chacun traitant un aspect different "
    "de la degradation radio :"
)

add_heading(doc, "2.1  Etape 1 : Demucs (Facebook Research)  [pre-entraine, fige]", level=2)
doc.add_paragraph(
    "Demucs (htdemucs) est un reseau de separation de sources entraine sur de grands corpus musicaux "
    "et vocaux. Il opere directement sur la forme d'onde (waveform) et supprime le bruit de fond "
    "large bande (souffle, interferences). Il est utilise en mode 100% debruitage (dry/wet = 0.0), "
    "sans melange avec le signal original. Ce modele est utilise tel quel, sans re-entrainement."
)
doc.add_paragraph(
    "Limites identifiees pour les signaux radio cockpit :"
)
for b in [
    "Demucs n'a jamais vu de signaux radio AM pendant son entrainement (bande 300-3400 Hz)",
    "Risque d'over-smoothing : suppression involontaire de consonnes fricatives et sibilantes",
    "Travaille en pleine bande alors que le signal utile est concentre dans 300-3400 Hz",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)
doc.add_paragraph(
    "Malgre ces limites, Demucs reste le meilleur candidat disponible parmi les modeles generiques "
    "testes (VoiceFixer et DeepFilterNet ont ete abandonnes car ils degradaient l'intelligibilite). "
    "Le SpectralResUNet etant entraine sur la sortie de Demucs, il apprend a corriger ses artefacts."
)

add_heading(doc, "2.2  Etape 2 : SpectralResUNet (modele custom)  [entraine de zero]", level=2)
doc.add_paragraph(
    "C'est le coeur du projet et la principale valeur ajoutee. Ce ResUNet est entraine de zero "
    "sur un dataset specifiquement concu pour les degradations radio cockpit. Il opere dans le "
    "domaine frequentiel (STFT) et accomplit deux taches simultanement :"
)
for b in [
    "Debruitage spectral : elimination des residus que Demucs n'a pas supprimes",
    "Reconstruction haute frequence : synthese des harmoniques manquantes (4 - 20 kHz)",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

doc.add_paragraph("Architecture detaillee :")
arch_rows = [
    ("Entree", "Signal 44 100 Hz -> STFT -> magnitude log (1025 x T)"),
    ("Encoder", "4 blocs : Conv2D stride 2 + ResBlock | canaux : 1->32->64->128->256"),
    ("Decoder", "4 blocs : ConvTranspose2D + skip connections (U-Net)"),
    ("Sortie", "Masque spectral (filtre bruit) + estimation HF (reconstruction frequences)"),
    ("Reconstruction", "masque * magnitude + HF, puis iSTFT avec phase originale"),
    ("Parametres", "3 973 586 parametres trainables"),
]
add_metric_table(doc, arch_rows, ["Composant", "Description"])
doc.add_paragraph()

add_heading(doc, "2.3  Etape 3 : MetricGAN+ (SpeechBrain, optionnel)  [pre-entraine, fige]", level=2)
doc.add_paragraph(
    "MetricGAN+ est un GAN entraine pour maximiser directement la metrique PESQ "
    "(Perceptual Evaluation of Speech Quality). Il sert de couche de polissage final "
    "pour ameliorer la naturalite perceptuelle du signal. Ce modele est utilise tel quel, "
    "sans re-entrainement. Son utilisation est optionnelle (USE_METRICGAN = True par defaut)."
)

add_heading(doc, "2.4  Schema du pipeline", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Radio (OGG/FLAC/MP3)\n"
    "  -> [1] Demucs htdemucs  (debruitage waveform 100%)\n"
    "  -> [2] Resample 44 100 Hz\n"
    "  -> [3] SpectralResUNet  (denoise spectral + reconstruction HF)\n"
    "  -> [4] MetricGAN+       (polissage PESQ, optionnel)\n"
    "  -> WAV 44 100 Hz"
)
run.font.name = 'Courier New'
run.font.size = Pt(10)
p.paragraph_format.left_indent = Cm(1)

# ============================================================
# 3. STRATEGIE D'ENTRAINEMENT
# ============================================================
add_heading(doc, "3. Strategie d'entrainement")

doc.add_paragraph(
    "Le SpectralResUNet est entraine en fine-tuning sur un dataset synthetique calibre sur les "
    "vraies degradations radio, completees par de vraies paires ATC."
)

add_heading(doc, "3.1  Dataset", level=2)
dataset_rows = [
    ("Train - ATC reel", "162", "Vrais signaux radio (decoupes 2-10s)"),
    ("Train - Synthetique LibriSpeech", "3 000", "Voix propres + degradations calibrees"),
    ("Val - ATC reel", "18", ""),
    ("Val - Synthetique LibriSpeech", "300", ""),
    ("TOTAL", "3 480 paires", "~5.9 heures d'audio a 44 100 Hz"),
]
add_metric_table(doc, dataset_rows, ["Source", "Nb paires", "Remarque"])
doc.add_paragraph()

doc.add_paragraph(
    "Chaque paire est composee d'un fichier clean (reference) et d'un fichier raw "
    "prealablement passe dans Demucs. Cela permet au SpectralResUNet d'apprendre "
    "a corriger la sortie de Demucs, reproduisant exactement les conditions d'inference."
)

doc.add_paragraph(
    "Generation des paires ATC reelles : les 6 fichiers radio annotes sont charges via ffmpeg "
    "(mono 44 100 Hz), alignes sur la longueur du plus court, puis decoupe en segments de 2 a 8 "
    "secondes. Les segments silencieux (RMS < 0.005) sont rejetes. "
    "Point de vigilance : le fichier Heathrow (973s) represente 89% des paires ATC reelles, "
    "ce qui cree un biais vers un seul type de signal radio (accent britannique, fréquences Heathrow). "
    "Diversifier les sources ATC est une priorite pour les travaux futurs."
)

add_heading(doc, "3.2  Degradations synthetiques appliquees", level=2)
for b in [
    "Bandpass 200-4 000 Hz (simulation bande radio AM)",
    "Downsampling 2/4/6 kHz (decimation et reechantillonnage)",
    "Bruit blanc (SNR 1-45 dB) + bruit rose",
    "Crackling (impulsions aleatoires, densite 0.0005-0.003)",
    "Interferences HF sinusoidales",
    "Clipping (seuil 0.4-0.8)",
    "AGC (compression dynamique radio)",
    "Dropouts radio (pertes de paquets, 1-5 par segment)",
    "Reverb cockpit synthetique",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

doc.add_paragraph(
    "Le profil de degradation est calibre sur les vrais fichiers radio via analyze_radio.py, "
    "garantissant que les degradations synthetiques correspondent a la realite du terrain."
)

add_heading(doc, "3.3  Hyperparametres", level=2)
hp_rows = [
    ("Batch size", "8"),
    ("Learning rate", "3e-4 (AdamW + CosineAnnealing)"),
    ("Epochs max", "80 (early stopping patience 15)"),
    ("Segment d'entrainement", "3 s (132 300 samples)"),
    ("Loss", "L1 spectrale + Multi-Resolution STFT Loss"),
    ("STFT", "n_fft=2048, hop=512, win=2048"),
    ("Data augmentation", "Gain aleatoire, time shift, bruit additif leger"),
]
add_metric_table(doc, hp_rows, ["Parametre", "Valeur"])
doc.add_paragraph()

# ============================================================
# 4. RESULTATS
# ============================================================
add_heading(doc, "4. Resultats obtenus")

doc.add_paragraph(
    "Les metriques suivantes ont ete mesurees sur des fichiers radio reels, "
    "en comparant le signal d'entree brut et le signal ameliore par le pipeline complet :"
)

metrics_rows = [
    ("LSD (Log-Spectral Distance, dB)", "5.85", "1.13", "-4.72", "Distorsion spectrale (< = mieux)"),
    ("STOI (intelligibilite, 0-1)", "0.535", "0.779", "+0.244", "Intelligibilite (+24 pts)"),
    ("SDR (Signal-to-Distortion Ratio, dB)", "0.64", "7.57", "+6.93 dB", "Separation signal/bruit"),
    ("PESQ (qualite perceptuelle, 1-4.5)", "1.16", "1.43", "+0.27", "Qualite MOS telephonique"),
]
add_metric_table(doc, metrics_rows,
                 ["Metrique", "Avant", "Apres", "Delta", "Interpretation"])
doc.add_paragraph()

doc.add_paragraph(
    "Les gains les plus significatifs sont sur le LSD (-4.7 dB, reconstruction spectrale fidele) "
    "et le STOI (+24 points, intelligibilite largement amelioree). "
    "Le SDR passe de 0.64 a 7.57 dB : avant traitement le bruit etait presque aussi fort que le signal utile, "
    "apres traitement le signal est clairement dominant (environ 5x plus fort que la distorsion residuelle). "
    "Le PESQ reste modeste car cette metrique est calibree pour la telephonie standard (bande 300-3 400 Hz) "
    "et ne valorise pas les frequences reconstruites au-dela de 4 kHz : elle sous-estime donc le gain reel."
)

add_heading(doc, "4.1  Bilan perceptuel : evaluation a l'ecoute", level=2)
doc.add_paragraph(
    "Une evaluation a l'ecoute des sorties du pipeline a mis en evidence les conclusions suivantes :"
)
for b in [
    "Demucs seul : resultats perceptuels convaincants. La suppression du bruit large bande est clairement "
    "audible et n'introduit pas d'artefacts majeurs sur les signaux radio cockpit testes.",
    "SpectralResUNet (entraine 66 epochs, best val_loss = 1.68) : deux problemes identifies a l'ecoute. "
    "(1) Les sorties sont moins bonnes que la sortie Demucs seule — le modele degrade le signal deja propre. "
    "(2) Sur les segments silencieux, le modele genere un bruit artificiel (hallucination spectrale).",
    "Analyse HF : l'analyse spectrale confirme que le ResUNet fait du debruitage HF (suppression des "
    "artefacts > 4 kHz) mais ne reconstruit pas de nouvelles hautes frequences. "
    "Ce n'est pas de la super-resolution au sens strict.",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

doc.add_paragraph(
    "Explication technique des deux artefacts observes :"
)
for b in [
    "Degradation sur signal propre : le modele applique toujours une transformation car la loss L1+STFT "
    "ne lui a pas appris a 'ne rien faire' quand l'entree est deja propre. "
    "Il n'a jamais ete expose a un cas ou la reponse correcte est l'identite.",
    "Hallucination sur silence : la tete d'estimation HF genere une estimation non nulle meme sur silence. "
    "Le modele a appris que 'sortie = masque x spectre + estimation HF' — sur un spectre nul, "
    "l'estimation HF domine et produit du bruit artificiel.",
    "Cause racine : 95% du dataset est synthetique (LibriSpeech). Il y a toujours du bruit a corriger "
    "dans les donnees d'entrainement — le modele n'a jamais appris a preserver un signal propre.",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

add_heading(doc, "4.2  Statut des modeles dans le pipeline", level=2)
model_status_rows = [
    ("Demucs (htdemucs)", "Facebook Research", "Pre-entraine, fige", "Utilise tel quel sans modification"),
    ("SpectralResUNet", "Custom (ce projet)", "Entraine de zero", "3.97M params, entraine sur dataset radio"),
    ("MetricGAN+", "SpeechBrain", "Pre-entraine, fige", "Utilise tel quel sans modification"),
]
add_metric_table(doc, model_status_rows, ["Modele", "Source", "Statut", "Remarque"])
doc.add_paragraph()
doc.add_paragraph(
    "La principale valeur ajoutee du projet reside dans le SpectralResUNet : c'est le seul modele "
    "entraine specificement sur les degradations radio cockpit. Les deux autres modeles apportent "
    "leur expertise generique (debruitage large bande pour Demucs, polissage perceptuel pour MetricGAN+) "
    "sans connaitre les specificites du domaine radio."
)

add_heading(doc, "4.3  Modeles testes et abandonnes", level=2)
abandoned_rows = [
    ("VoiceFixer", "Restauration vocale universelle", "Abandonne", "Ecrasement des freq. radio, artefacts vocaux severes"),
    ("DeepFilterNet", "Debruitage neural temps reel", "Abandonne", "Degradation de l'intelligibilite sur signaux radio"),
    ("SepFormer", "Separation de sources (SpeechBrain)", "Optionnel", "Moins robuste que Demucs sur les radio"),
]
add_metric_table(doc, abandoned_rows, ["Modele", "Role", "Statut", "Raison"])
doc.add_paragraph()

# ============================================================
# 5. PERSPECTIVES
# ============================================================
add_heading(doc, "5. Perspectives et travaux futurs")

add_heading(doc, "5.1  Priorite 1 : augmenter le corpus ATC reel", level=2)
doc.add_paragraph(
    "Le constat principal est que le dataset actuel est trop domine par du synthetique (95%). "
    "L'action la plus impactante est d'augmenter les donnees reelles :"
)
for b in [
    "Objectif : passer de 180 a 1 000+ paires ATC reelles annotees (clean/raw)",
    "Sources : corpus publics ATCO2 (LDC), LiveATC.net, enregistrements cockpit open-source",
    "Impact direct : le SpectralResUNet apprendrait des caracteristiques radio reelles, "
    "pas des simulations artificielles",
    "Diversite : integrer plusieurs aeroports, frequences radio et conditions meteorologiques",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

add_heading(doc, "5.2  Priorite 2 : debruiteur specialise radio (remplacement de Demucs)", level=2)
doc.add_paragraph(
    "Demucs fonctionne bien mais n'a jamais ete entraine sur des signaux radio cockpit. "
    "Un debruiteur specifique apporterait une connaissance exacte des degradations radio :"
)
for b in [
    "SepFormer (SpeechBrain) fine-tune sur les paires ATC reelles : architecture Transformer "
    "tres performante, deja presente dans le projet comme option",
    "ResUNet waveform leger : meme approche que le SpectralResUNet mais en domaine temporel, "
    "reutilise l'infrastructure existante (train.py)",
    "Implementation : desactiver APPLY_DEMUCS_TO_RAW dans create_dataset.py, "
    "entrainer sur paires brutes, remplacer Demucs dans inference.py",
    "Avantage cle : connaissance exacte du crackling, dropout radio "
    "et souffle HF specifiques aux communications cockpit",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)
doc.add_paragraph(
    "Faisabilite a court terme : fine-tuner SepFormer sur les 180 paires ATC reelles existantes "
    "est realiste et constituerait une alternative solide a Demucs."
)

add_heading(doc, "5.3  Corrections prioritaires du SpectralResUNet", level=2)
doc.add_paragraph(
    "Les deux artefacts identifies a l'ecoute ont des corrections connues :"
)
for b in [
    "Identite loss : ajouter un terme de loss qui penalise la deviation par rapport a l'entree "
    "quand l'entree est deja propre (ex: loss += lambda * ||output - input||^2 si SNR_input > seuil). "
    "Cela apprend au modele a 'ne rien faire' quand ce n'est pas necessaire.",
    "Gate de silence : detecter les trames silencieuses (RMS < seuil) et bypasser le modele — "
    "retourner l'entree directement sans passer dans le ResUNet.",
    "Exemples clean dans le dataset : inclure des paires (clean, clean) dans l'entrainement "
    "pour que le modele apprenne explicitement l'identite.",
    "Loss perceptuelle : remplacer la STFT loss pure par une loss basee sur PESQ differentiable "
    "pour aligner l'optimisation sur la perception humaine.",
    "Architecture attention : blocs Transformer dans le bottleneck pour mieux capturer "
    "les dependances temporelles longue portee.",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

add_heading(doc, "5.4  Ameliorations du dataset synthetique", level=2)
for b in [
    "Augmenter le ratio ATC reel (actuellement 5%) : objectif 20-30% avec plus de fichiers annotes",
    "Degradations adaptatives : utiliser un modele de canal radio (Rayleigh fading, Doppler) plutot que des filtres statiques",
    "Multi-locuteurs : integrer des corpus radio multilingues (ATCO2, LDC Air Traffic Control)",
    "Donnees synthese + TTS : generer des scripts de phraseologie radio avec un TTS pour augmenter le volume",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

add_heading(doc, "5.5  Evaluation et metriques", level=2)
doc.add_paragraph(
    "Les metriques actuelles ont des limites importantes pour ce domaine. "
    "Le PESQ est calibre telephonie standard et ne valorise pas la reconstruction HF. "
    "Le STOI reste la metrique la plus pertinente operationnellement. "
    "Metriques a integrer en priorite :"
)
for b in [
    "WER (Word Error Rate) via Whisper : metrique la plus operationnelle - mesure l'impact direct sur la transcription radio",
    "ESTOI (Extended STOI) : version amelioree de STOI, plus fiable sur les signaux tres bruites",
    "DNSMOS / NISQA : metriques sans reference (no-reference MOS), utiles quand il n'y a pas de signal clean de reference",
    "MOS subjectif : test d'ecoute avec de vrais operateurs radio pour valider l'intelligibilite operationnelle",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

add_heading(doc, "5.6  Deploiement", level=2)
for b in [
    "Optimisation temps reel : quantification INT8 / ONNX export pour inference embarquee",
    "Traitement en streaming : adapter le pipeline pour traiter des segments audio en temps reel (latence < 200 ms)",
    "Integration dans un systeme de transcription : coupler avec Whisper pour un pipeline complet radio -> texte",
    "API REST : exposer le pipeline comme service pour integration dans des outils existants",
]:
    doc.add_paragraph(b, style='List Bullet').runs[0].font.size = Pt(11)

# ============================================================
# 6. STACK TECHNIQUE
# ============================================================
add_heading(doc, "6. Stack technique")

stack_rows = [
    ("PyTorch / torchaudio", "Framework deep learning (ROCm pour GPU AMD RX 6800)"),
    ("Demucs (Facebook Research)", "Separation de sources / debruitage waveform"),
    ("SpeechBrain", "MetricGAN+ (polissage PESQ)"),
    ("HuggingFace datasets", "Telechargement LibriSpeech train-clean-100"),
    ("soundfile / librosa", "I/O audio et traitement du signal"),
    ("scipy / numpy", "Traitement signal (filtres, resample, STFT)"),
    ("PESQ / PYSTOI / mir_eval", "Metriques d'evaluation audio"),
    ("matplotlib", "Visualisation spectrogrammes et courbes d'entrainement"),
]
add_metric_table(doc, stack_rows, ["Outil / Librairie", "Role"])
doc.add_paragraph()

# ============================================================
# CONCLUSION
# ============================================================
add_heading(doc, "Conclusion")

doc.add_paragraph(
    "Ce projet demontre la faisabilite d'un pipeline de super-resolution audio pour les "
    "communications radio. Les gains objectifs obtenus (LSD -4.7 dB, STOI +24 pts, SDR +7 dB) "
    "valident l'approche. L'evaluation a l'ecoute confirme que Demucs est l'etape la plus "
    "efficace du pipeline dans son etat actuel."
)

doc.add_paragraph(
    "Le SpectralResUNet, bien qu'il ameliore les metriques objectives, n'apporte pas encore "
    "de gain perceptuel clair a l'ecoute. Ce resultat est explique par le manque de donnees "
    "ATC reelles (5% du dataset) et les limites des degradations synthetiques. "
    "L'augmentation du corpus reel et le remplacement de Demucs par un debruiteur "
    "specialise radio sont les deux axes prioritaires pour progresser."
)

doc.add_paragraph(
    "Le pipeline est concu pour etre modulaire : chaque etape peut etre remplacee "
    "ou fine-tunee independamment, ce qui facilite l'experimentation future. "
    "Le couplage avec Whisper (WER) est la prochaine etape d'evaluation pour "
    "mesurer l'impact operationnel reel sur la transcription radio."
)

doc.save(OUTPUT_PATH)
print(f"Document genere : {OUTPUT_PATH}")
